from __future__ import annotations

import logging
from dataclasses import dataclass
from io import BytesIO

from app.connectors.llm import EmbeddingProviderError
from app.connectors.storage import StorageError
from app.knowledge.models import KnowledgeChunk, KnowledgeDocument, KnowledgeDocumentStatus
from app.knowledge.vector_store import EMBEDDING_DIM, VectorChunk

logger = logging.getLogger(__name__)

CHUNK_TOKENS = 512
OVERLAP_TOKENS = 64
CHARS_PER_TOKEN = 4
EXCERPT_CHARS = 200
EMBED_BATCH_SIZE = 32
MAX_FILENAME_CHARS = 120

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


class EmbeddingNotConfiguredError(RuntimeError):
    """Raised when ingestion has no embedding connector to encode chunks."""


class UnsupportedContentTypeError(ValueError):
    """Raised when a document content type has no registered parser."""


class DocumentParseError(ValueError):
    """Raised when a document parser cannot extract content from the bytes."""


class NoExtractableTextError(ValueError):
    """Raised when a document parses successfully but yields no text."""


@dataclass
class ParsedSection:
    text: str
    page_or_sheet: str = ""


@dataclass
class ChunkDraft:
    content: str
    token_estimate: int
    page_or_sheet: str
    excerpt: str


def sanitize_filename(filename: str) -> str:
    name = (filename or "").replace("\\", "/").rsplit("/", 1)[-1].strip()
    name = " ".join(name.split())
    if len(name) > MAX_FILENAME_CHARS:
        name = name[:MAX_FILENAME_CHARS].rstrip()
    return name or "document"


def _failure_message(error: Exception) -> str:
    if isinstance(error, EmbeddingNotConfiguredError):
        return "embedding provider is not configured"
    if isinstance(error, EmbeddingProviderError):
        return "embedding provider request failed"
    if isinstance(error, StorageError):
        return "file storage failed"
    if isinstance(error, UnsupportedContentTypeError):
        return "unsupported file type"
    if isinstance(error, DocumentParseError):
        return "document could not be parsed"
    if isinstance(error, NoExtractableTextError):
        return "no extractable text found in document"
    return "ingestion failed unexpectedly"


def parse_pdf(content: bytes) -> list[ParsedSection]:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    sections: list[ParsedSection] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            sections.append(ParsedSection(text=text, page_or_sheet=f"page {index}"))
    return sections


def parse_docx(content: bytes) -> list[ParsedSection]:
    from docx import Document

    document = Document(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraph for paragraph in paragraphs if paragraph)
    if not text:
        return []
    return [ParsedSection(text=text, page_or_sheet="")]


def parse_xlsx(content: bytes) -> list[ParsedSection]:
    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    sections: list[ParsedSection] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                rows.append(" | ".join(cells))
        text = "\n".join(rows)
        if text:
            sections.append(ParsedSection(text=text, page_or_sheet=sheet.title))
    return sections


CONTENT_TYPE_PARSERS = {
    PDF_CONTENT_TYPE: parse_pdf,
    DOCX_CONTENT_TYPE: parse_docx,
    XLSX_CONTENT_TYPE: parse_xlsx,
}


def chunk_sections(
    sections: list[ParsedSection],
    chunk_size: int = CHUNK_TOKENS,
    overlap: int = OVERLAP_TOKENS,
) -> list[ChunkDraft]:
    chunk_chars = max(1, chunk_size * CHARS_PER_TOKEN)
    overlap_chars = max(0, overlap * CHARS_PER_TOKEN)
    chunks: list[ChunkDraft] = []
    for section in sections:
        text = section.text.strip()
        if not text:
            continue
        start = 0
        while start < len(text):
            piece = text[start : start + chunk_chars]
            chunks.append(
                ChunkDraft(
                    content=piece,
                    token_estimate=max(1, len(piece) // CHARS_PER_TOKEN),
                    page_or_sheet=section.page_or_sheet,
                    excerpt=piece[:EXCERPT_CHARS],
                )
            )
            if start + chunk_chars >= len(text):
                break
            start += chunk_chars - overlap_chars
    return chunks


class IngestionService:
    def __init__(self, session, *, storage, embedding=None, vector_store=None) -> None:
        self.session = session
        self.storage = storage
        self.embedding = embedding
        self.vector_store = vector_store

    def create_document(
        self,
        *,
        organization_id: str,
        filename: str,
        content_type: str,
        size: int,
        product_line_id: str | None = None,
    ) -> KnowledgeDocument:
        document = KnowledgeDocument(
            organization_id=organization_id,
            filename=sanitize_filename(filename),
            content_type=content_type,
            size=size,
            product_line_id=product_line_id,
        )
        self.session.add(document)
        self.session.flush()
        return document

    async def process(self, document_id: str, organization_id: str, content: bytes) -> KnowledgeDocument:
        document = self.session.get(KnowledgeDocument, document_id)
        if document is None or document.organization_id != organization_id:
            raise LookupError("knowledge document not found")
        document.status = KnowledgeDocumentStatus.PROCESSING
        document.failure_message = ""
        self.session.commit()
        storage_key = f"{document.organization_id}/{document.id}/{document.filename}"
        try:
            # Ingestion is intentionally synchronous for V1: the request blocks until the
            # document is READY or FAILED. The PROCESSING status is retained as the seam for
            # a future background worker, which would persist it and offload this work.
            await self._ingest(document, content, storage_key)
        except Exception as error:
            logger.exception("knowledge document ingestion failed: %s", document_id)
            self.session.rollback()
            await self._delete_stored_object(storage_key)
            document = self.session.get(KnowledgeDocument, document_id)
            document.status = KnowledgeDocumentStatus.FAILED
            document.failure_message = _failure_message(error)
            self.session.commit()
        return document

    async def _ingest(self, document: KnowledgeDocument, content: bytes, storage_key: str) -> None:
        parser = CONTENT_TYPE_PARSERS.get(document.content_type)
        if parser is None:
            raise UnsupportedContentTypeError(document.content_type)
        if self.embedding is None:
            raise EmbeddingNotConfiguredError()
        try:
            sections = parser(content)
        except Exception as error:
            raise DocumentParseError(document.content_type) from error
        chunks = chunk_sections(sections)
        if not chunks:
            raise NoExtractableTextError()
        embeddings = await self._embed_chunks([chunk.content for chunk in chunks])

        # Store the original file first; the failure handler best-effort deletes it if
        # anything after this point fails, before persisting the FAILED status.
        await self.storage.put(storage_key, content)

        for index, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            persisted = KnowledgeChunk(
                document_id=document.id,
                organization_id=document.organization_id,
                chunk_index=index,
                content=chunk.content,
                token_estimate=chunk.token_estimate,
                page_or_sheet=chunk.page_or_sheet,
                excerpt=chunk.excerpt,
            )
            self.session.add(persisted)
            self.session.flush()
            if self.vector_store is not None:
                await self.vector_store.upsert(
                    VectorChunk(
                        chunk_id=persisted.id,
                        organization_id=document.organization_id,
                        document_id=document.id,
                        product_line_id=document.product_line_id,
                        embedding=embedding,
                    )
                )
        document.status = KnowledgeDocumentStatus.READY
        document.failure_message = ""
        self.session.commit()

    async def _embed_chunks(self, texts: list[str]) -> list[list[float]]:
        vectors: list[list[float]] = []
        for start in range(0, len(texts), EMBED_BATCH_SIZE):
            batch = texts[start : start + EMBED_BATCH_SIZE]
            for vector in await self.embedding.embed(batch):
                if len(vector) != EMBEDDING_DIM:
                    raise EmbeddingProviderError(
                        f"embedding provider returned dimension {len(vector)}, expected {EMBEDDING_DIM}"
                    )
                vectors.append(vector)
        if len(vectors) != len(texts):
            raise EmbeddingProviderError("embedding provider returned an unexpected number of vectors")
        return vectors

    async def _delete_stored_object(self, key: str) -> None:
        try:
            await self.storage.delete(key)
        except Exception:
            logger.exception("failed to delete stored object after ingestion failure: %s", key)
